"""
구조화된 회의록 데이터를 DOCX / PDF / HWP 파일로 생성
"""
from __future__ import annotations

import os
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console

console = Console()

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
OUTPUT_DIR = Path(__file__).parent.parent / "output"


# ──────────────────────────────────────────────
# DOCX 생성
# ──────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table_row(table, key: str, value: str):
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True


def generate_docx(meeting_data: dict, output_path: Path) -> Path:
    doc = Document()

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # 제목
    title = doc.add_heading("", level=0)
    run = title.add_run("회  의  록")
    run.font.size = Pt(20)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 기본 정보 테이블
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.0)

    _add_table_row(table, "회의명", meeting_data.get("회의명", ""))
    _add_table_row(table, "일  시", meeting_data.get("일시", ""))
    _add_table_row(table, "장  소", meeting_data.get("장소", ""))
    _add_table_row(table, "참석자", ", ".join(meeting_data.get("참석자", [])))

    # 안건
    agenda = meeting_data.get("안건", [])
    if agenda:
        agenda_text = "\n".join(f"{i+1}. {item}" for i, item in enumerate(agenda))
        _add_table_row(table, "안  건", agenda_text)

    doc.add_paragraph()

    # 주요 논의 내용
    discussions = meeting_data.get("주요_논의내용", [])
    if discussions:
        _add_heading(doc, "1. 주요 논의 내용", level=1)
        for i, item in enumerate(discussions):
            _add_heading(doc, f"  {i+1}) {item.get('안건', '')}", level=2)
            p = doc.add_paragraph(item.get("내용", ""))
            p.paragraph_format.left_indent = Cm(1.0)

    # 결정 사항
    decisions = meeting_data.get("결정사항", [])
    if decisions:
        _add_heading(doc, "2. 결정 사항", level=1)
        for item in decisions:
            p = doc.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Cm(0.5)

    # 액션 아이템 (내용이 있는 항목만)
    actions = [a for a in meeting_data.get("액션아이템", []) if a.get("내용", "").strip()]
    if actions:
        _add_heading(doc, "3. 액션 아이템", level=1)

        # 담당자/기한 컬럼이 하나라도 채워진 경우에만 테이블, 아니면 목록
        has_meta = any(a.get("담당자", "").strip() or a.get("기한", "").strip() for a in actions)
        if has_meta:
            action_table = doc.add_table(rows=1, cols=3)
            action_table.style = "Table Grid"
            hdr = action_table.rows[0].cells
            hdr[0].text = "담당자"
            hdr[1].text = "내용"
            hdr[2].text = "기한"
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True
            for action in actions:
                row = action_table.add_row()
                row.cells[0].text = action.get("담당자", "")
                row.cells[1].text = action.get("내용", "")
                row.cells[2].text = action.get("기한", "")
        else:
            for action in actions:
                p = doc.add_paragraph(f"• {action.get('내용', '')}")
                p.paragraph_format.left_indent = Cm(0.5)

    # 다음 회의 / 특이사항
    next_meeting = meeting_data.get("다음회의", "")
    if next_meeting:
        _add_heading(doc, "4. 다음 회의 일정", level=1)
        doc.add_paragraph(next_meeting)

    notes = meeting_data.get("특이사항", "")
    if notes:
        _add_heading(doc, "5. 특이 사항", level=1)
        doc.add_paragraph(notes)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    console.print(f"[green]DOCX 저장: {output_path}[/green]")
    return output_path


# ──────────────────────────────────────────────
# PDF 생성 (DOCX → PDF 변환)
# ──────────────────────────────────────────────

def generate_pdf(meeting_data: dict, output_path: Path) -> Path:
    from docx2pdf import convert

    # 임시 DOCX 생성 후 PDF 변환
    temp_docx = output_path.with_suffix(".docx")
    generate_docx(meeting_data, temp_docx)

    convert(str(temp_docx), str(output_path))
    temp_docx.unlink(missing_ok=True)

    console.print(f"[green]PDF 저장: {output_path}[/green]")
    return output_path


# ──────────────────────────────────────────────
# HWP 생성 (win32com + 한컴 HWP)
# ──────────────────────────────────────────────

def generate_hwp(meeting_data: dict, output_path: Path) -> Path:
    """
    템플릿 HWP를 복사한 뒤, win32com을 통해 플레이스홀더를 치환합니다.
    템플릿 HWP에는 아래 텍스트가 포함되어 있어야 합니다:
      {{회의명}}, {{일시}}, {{장소}}, {{참석자}},
      {{안건}}, {{논의내용}}, {{결정사항}}, {{액션아이템}}, {{다음회의}}
    """
    try:
        import win32com.client as win32
    except ImportError:
        console.print("[red]pywin32가 설치되어 있지 않습니다.[/red]")
        raise

    template_path = TEMPLATES_DIR / "회의록_template.hwp"
    if not template_path.exists():
        raise FileNotFoundError(f"HWP 템플릿이 없습니다: {template_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    # 플레이스홀더 → 실제 내용 매핑
    replacements = {
        "{{회의명}}": meeting_data.get("회의명", ""),
        "{{일시}}": meeting_data.get("일시", ""),
        "{{장소}}": meeting_data.get("장소", ""),
        "{{참석자}}": ", ".join(meeting_data.get("참석자", [])),
        "{{안건}}": "\n".join(
            f"{i+1}. {a}" for i, a in enumerate(meeting_data.get("안건", []))
        ),
        "{{논의내용}}": "\n".join(
            f"[{d.get('안건','')}]\n{d.get('내용','')}"
            for d in meeting_data.get("주요_논의내용", [])
        ),
        "{{결정사항}}": "\n".join(
            f"• {d}" for d in meeting_data.get("결정사항", [])
        ),
        "{{액션아이템}}": "\n".join(
            f"[{a.get('담당자','')}] {a.get('내용','')} ({a.get('기한','')})"
            for a in meeting_data.get("액션아이템", [])
        ),
        "{{다음회의}}": meeting_data.get("다음회의", ""),
        "{{특이사항}}": meeting_data.get("특이사항", ""),
    }

    hwp = None
    try:
        hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        hwp.Open(str(output_path.resolve()), "", "forceopen:true")

        for placeholder, value in replacements.items():
            action = hwp.CreateAction("AllReplace")
            param_set = action.CreateSet()
            action.GetDefault(param_set)
            param_set.SetItem("FindString", placeholder)
            param_set.SetItem("ReplaceString", value)
            param_set.SetItem("IgnoreMessage", 1)
            param_set.SetItem("FindType", 1)
            action.Execute(param_set)

        hwp.Save()
        console.print(f"[green]HWP 저장: {output_path}[/green]")

    finally:
        if hwp is not None:
            try:
                hwp.Quit()
            except Exception:
                pass

    return output_path


# ──────────────────────────────────────────────
# 통합 생성 함수
# ──────────────────────────────────────────────

GENERATORS = {
    "docx": generate_docx,
    "pdf": generate_pdf,
    "hwp": generate_hwp,
}


def generate_document(
    meeting_data: dict,
    output_format: str,
    output_filename: str | None = None,
) -> Path:
    """
    meeting_data: summarizer가 반환한 딕셔너리
    output_format: 'docx' | 'pdf' | 'hwp'
    output_filename: 저장 파일명 (확장자 제외). None이면 회의명+날짜로 자동 생성
    """
    fmt = output_format.lower().strip(".")
    if fmt not in GENERATORS:
        raise ValueError(f"지원하지 않는 형식: {fmt}. 사용 가능: {list(GENERATORS.keys())}")

    if not output_filename:
        title = meeting_data.get("회의명", "회의록").replace(" ", "_")
        date_str = datetime.now().strftime("%Y%m%d_%H%M")
        output_filename = f"{title}_{date_str}"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{output_filename}.{fmt}"

    return GENERATORS[fmt](meeting_data, output_path)
